import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def _build_output_path(filename: str) -> str:
    """Build the default DOCX output path for a generated article."""

    return os.path.join("output", "articles", f"{filename}.docx")


def _fallback_output_path(filename: str) -> str:
    """Build a timestamped fallback path when the default file is locked."""

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return os.path.join("output", "articles", f"{filename}_{timestamp}.docx")

def save_article(article_text: str, filename: str, image_path: str = None) -> str:
    """
    Saves generated article text into a professionally formatted DOCX file.

    Features:
    - Professional typography and styling
    - Colored headings with proper hierarchy
    - Featured image placement
    - Bullet point formatting
    - FAQ formatting with bold questions
    - Clean paragraph spacing
    """

    path = _build_output_path(filename)
    os.makedirs(os.path.dirname(path), exist_ok=True)

    doc = Document()

    # -----------------------------
    # Define Custom Styles
    # -----------------------------
    styles = doc.styles

    # Main Title Style (H1)
    title_style = styles.add_style('ArticleTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Georgia'
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(30, 30, 30)
    title_style.paragraph_format.space_after = Pt(6)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section Heading Style (H2)
    h2_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(18)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(102, 126, 234)  # Purple-blue gradient color
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)

    # Subsection Heading Style (H3)
    h3_style = styles.add_style('SubsectionHeading', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.name = 'Arial'
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(80, 80, 80)
    h3_style.paragraph_format.space_before = Pt(14)
    h3_style.paragraph_format.space_after = Pt(6)

    # Body Text Style
    body_style = styles.add_style('ArticleBody', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = 'Georgia'
    body_style.font.size = Pt(11)
    body_style.font.color.rgb = RGBColor(50, 50, 50)
    body_style.paragraph_format.space_after = Pt(8)
    body_style.paragraph_format.line_spacing = 1.4

    # Bullet Point Style
    bullet_style = styles.add_style('ArticleBullet', WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.font.name = 'Georgia'
    bullet_style.font.size = Pt(11)
    bullet_style.font.color.rgb = RGBColor(50, 50, 50)
    bullet_style.paragraph_format.space_after = Pt(4)
    bullet_style.paragraph_format.left_indent = Inches(0.25)

    # FAQ Question Style
    faq_q_style = styles.add_style('FAQQuestion', WD_STYLE_TYPE.PARAGRAPH)
    faq_q_style.font.name = 'Arial'
    faq_q_style.font.size = Pt(11)
    faq_q_style.font.bold = True
    faq_q_style.font.color.rgb = RGBColor(102, 126, 234)
    faq_q_style.paragraph_format.space_before = Pt(12)
    faq_q_style.paragraph_format.space_after = Pt(4)

    # FAQ Answer Style
    faq_a_style = styles.add_style('FAQAnswer', WD_STYLE_TYPE.PARAGRAPH)
    faq_a_style.font.name = 'Georgia'
    faq_a_style.font.size = Pt(11)
    faq_a_style.font.color.rgb = RGBColor(60, 60, 60)
    faq_a_style.paragraph_format.space_after = Pt(8)
    faq_a_style.paragraph_format.left_indent = Inches(0.15)

    # -----------------------------
    # Document Title
    # -----------------------------
    title = filename.replace("_", " ").title()
    title_para = doc.add_paragraph(title, style='ArticleTitle')

    # Add a subtle line under title
    doc.add_paragraph("─" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -----------------------------
    # Featured Image
    # -----------------------------
    if image_path and os.path.exists(image_path):
        doc.add_paragraph()  # spacing before image

        # Add image centered
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(image_path, width=Inches(5.5))

        # Image caption
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run(f"Featured Image: {title}")
        caption_run.font.size = Pt(9)
        caption_run.font.italic = True
        caption_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # spacing after image

    # -----------------------------
    # Body Formatting
    # -----------------------------
    lines = article_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # H2 - Main Section Heading (## )
        if line.startswith("## "):
            heading_text = line.replace("## ", "").strip()
            doc.add_paragraph(heading_text, style='SectionHeading')
            i += 1
            continue

        # H3 - Subsection Heading (### )
        if line.startswith("### "):
            heading_text = line.replace("### ", "").strip()
            doc.add_paragraph(heading_text, style='SubsectionHeading')
            i += 1
            continue

        # FAQ Question (**Q: ... **)
        if line.startswith("**Q:") or line.startswith("**Q "):
            # Remove markdown bold markers
            question_text = re.sub(r'\*\*', '', line)
            doc.add_paragraph(question_text, style='FAQQuestion')
            i += 1
            continue

        # FAQ Answer (A: ...)
        if line.startswith("A:") or line.startswith("**A:"):
            answer_text = re.sub(r'\*\*', '', line)
            doc.add_paragraph(answer_text, style='FAQAnswer')
            i += 1
            continue

        # Bullet points (- or • )
        if line.startswith(("- ", "• ", "* ")):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(style='ArticleBullet')

            # Add bullet character
            bullet_run = p.add_run("● ")
            bullet_run.font.color.rgb = RGBColor(102, 126, 234)
            bullet_run.font.size = Pt(10)

            # Add bullet text - handle bold text within bullets
            if "**" in bullet_text:
                parts = re.split(r'(\*\*[^*]+\*\*)', bullet_text)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        bold_run = p.add_run(part[2:-2])
                        bold_run.font.bold = True
                        bold_run.font.name = 'Georgia'
                        bold_run.font.size = Pt(11)
                    elif part:
                        normal_run = p.add_run(part)
                        normal_run.font.name = 'Georgia'
                        normal_run.font.size = Pt(11)
            else:
                text_run = p.add_run(bullet_text)
                text_run.font.name = 'Georgia'
                text_run.font.size = Pt(11)

            i += 1
            continue

        # Normal paragraph - handle inline bold/italic
        p = doc.add_paragraph(style='ArticleBody')

        # Process inline formatting (bold with **)
        if "**" in line:
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    bold_run = p.add_run(part[2:-2])
                    bold_run.font.bold = True
                    bold_run.font.name = 'Georgia'
                    bold_run.font.size = Pt(11)
                    bold_run.font.color.rgb = RGBColor(50, 50, 50)
                elif part:
                    normal_run = p.add_run(part)
                    normal_run.font.name = 'Georgia'
                    normal_run.font.size = Pt(11)
                    normal_run.font.color.rgb = RGBColor(50, 50, 50)
        else:
            run = p.add_run(line)
            run.font.name = 'Georgia'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(50, 50, 50)

        i += 1

    # -----------------------------
    # Footer separator
    # -----------------------------
    doc.add_paragraph()
    footer_line = doc.add_paragraph("─" * 50)
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer text
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Generated by AI Article Generator")
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(150, 150, 150)

    # -----------------------------
    # Save Document
    # -----------------------------
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback_path = _fallback_output_path(filename)
        doc.save(fallback_path)
        return fallback_path
